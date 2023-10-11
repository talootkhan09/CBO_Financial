from django.shortcuts import render, redirect,get_object_or_404
from django.http import HttpResponse, JsonResponse
from django.core.files.storage import FileSystemStorage
import os
import pickle
import openai
import logging
from PyPDF2 import PdfReader
from pdfquery import PDFQuery
from pdfquery.cache import FileCache
from pdf2image import convert_from_path
from pdf2image.exceptions import PDFPageCountError
from django.conf import settings
from django.template import loader
from django.views.decorators.csrf import csrf_exempt
from docx import Document  # Added for DOCX support
from openpyxl import load_workbook  # Added for XLSX support
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores.faiss import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.callbacks import get_openai_callback
from langchain.llms import OpenAI
from langchain.chat_models import ChatOpenAI
from .models import Client, FundingProgram, Document
from .forms import ClientForm, FundingProgramForm, DocumentForm

def upload_file(request):
    clients = Client.objects.all()
    if request.method == "POST" and request.FILES["uploaded_file"]:
        uploaded_file = request.FILES["uploaded_file"]
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        # Check the file extension and handle accordingly
        if file_extension in ['pdf', 'docx', 'xlsx']:
            # Save the uploaded file to a temporary location
            fs = FileSystemStorage()
            file_path = fs.save(uploaded_file.name, uploaded_file)

            # Process the uploaded file based on its extension
            if file_extension == 'pdf':
                text = extract_text_from_pdf(file_path)
            elif file_extension == 'docx':
                text = extract_text_from_docx(file_path)
            elif file_extension == 'xlsx':
                text = extract_text_from_xlsx(file_path)

            client_id = request.POST.get('client')

            # Create a new Document instance and associate it with the client and program
            document = Document(
                client=Client.objects.get(id=client_id),
                uploaded_file=file_path  # Save the path to the file
            )
            document.save()
            # Calculate embeddings and save them
            store_name = uploaded_file.name[:-4]
            vector_store_path = os.path.join(
                settings.MEDIA_ROOT, "vector_store", store_name
            )

            if not os.path.exists(vector_store_path):
                os.makedirs(vector_store_path)  # Create the directory if it doesn't exist

            if not os.path.exists(f"{vector_store_path}.pkl"):
                embeddings = OpenAIEmbeddings()

                # Ensure 'text' is a list of texts/documents
                text_list = [text]  # Assuming 'text' is a single document
                VectorStore = FAISS.from_texts(text_list, embedding=embeddings)
                with open(f"{vector_store_path}.pkl", "wb") as f:
                    pickle.dump(VectorStore, f)
                logging.info(f"Saved {vector_store_path}.pkl")

            # Render a template with a form for entering queries
            context = {"file_path": file_path}
            return render(request, "upload.html", context)

    return render(request, "upload.html",{'clients': clients})

@csrf_exempt
def run_query(request):
    if request.method == "POST":
        file_path = request.POST.get("file_path")
        query = request.POST.get("query")
        if file_path and query:
            # Load the pre-computed embeddings
            store_name = os.path.basename(file_path)[:-4]
            vector_store_path = os.path.join(
                settings.MEDIA_ROOT, "vector_store", store_name
            )

            with open(f"{vector_store_path}.pkl", "rb") as f:
                VectorStore = pickle.load(f)
            # Perform the query and generate a response
            docs = VectorStore.similarity_search(query=query, k=3)

            llm = ChatOpenAI(model_name="gpt-3.5-turbo")
            chain = load_qa_chain(llm=llm, chain_type="stuff")

            with get_openai_callback() as cb:
                response = chain.run(input_documents=docs, question=query)

            # Return the response as JSON
            return JsonResponse({"response": response})

    return HttpResponse(status=400)

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        pdf_reader = PdfReader(pdf_path)
        num_pages = len(pdf_reader.pages)

        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + '\n'
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")

    return text

def extract_text_from_docx(docx_path):
    text = ""
    try:
        doc = Document(docx_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")

    return text

def extract_text_from_xlsx(xlsx_path):
    text = ""
    try:
        wb = load_workbook(filename=xlsx_path, read_only=True)
        for sheet in wb:
            for row in sheet.iter_rows(values_only=True):
                text += ' '.join(map(str, row)) + '\n'
    except Exception as e:
        print(f"Error extracting text from XLSX: {str(e)}")

    return text


def client_list(request):
    clients = Client.objects.all()
    return render(request, 'client_list.html', {'clients': clients})

def add_client(request):
    if request.method == 'POST':
        form = ClientForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('client_list')
    else:
        form = ClientForm()
    return render(request, 'add_client.html', {'form': form})

def edit_client(request, client_id):
    client = Client.objects.get(pk=client_id)
    if request.method == 'POST':
        form = ClientForm(request.POST, instance=client)
        if form.is_valid():
            form.save()
            return redirect('client_list')
    else:
        form = ClientForm(instance=client)
    return render(request, 'edit_client.html', {'form': form, 'client': client})

def delete_client(request, client_id):
    client = Client.objects.get(pk=client_id)
    if request.method == 'POST':
        client.delete()
        return redirect('client_list')
    return render(request, 'delete_client.html', {'client': client})

def document_list(request):
    documents = Document.objects.all()
    return render(request, "document_list.html", {"documents": documents})

def document_detail(request, document_id):
    document = get_object_or_404(Document, pk=document_id)
    return render(request, "document_detail.html", {"document": document})